# -*- coding: utf-8 -*-
"""
Generate the final Word version of the thesis from docs/毕业论文初稿.md.

The output follows the supplied undergraduate thesis template:
- A4, 3 cm margins, 2 cm header/footer distance.
- Body: 宋体 小四, 1.25 line spacing, first-line indent 2 Chinese chars.
- Chapter headings: 黑体 小三, centered; section headings: 黑体 四号/小四.
- Cover, Chinese abstract, English abstract, TOC, body, references and appendices.

Run from the repository root:
    python docs/md_to_thesis_docx.py
"""

from __future__ import annotations

import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "毕业论文初稿.md"
OUT_DOCX = Path(os.environ.get("THESIS_OUT_DOCX", ROOT / "毕业论文初稿.docx"))

BODY_FONT = "宋体"
LATIN_FONT = "Times New Roman"
HEADING_FONT = "黑体"
CODE_FONT = "Consolas"
SCHOOL_HEADER = "中国石油大学（北京）本科毕业设计（论文）"
ENGLISH_TITLE = "Design and Implementation of an Automated Testing Platform Based on Flaky Analysis and Adaptive Execution Strategy"


def normalize_figure_table_numbers(text: str) -> str:
    """Use the template's chapter-based figure/table number form, e.g. 图4.1 / 表6.1."""
    text = re.sub(r"([图表])\s*([0-9A-Z])-(\d+[a-z]?)", r"\1\2.\3", text)
    text = re.sub(r"([图表])\s+([0-9A-Z])\.(\d+[a-z]?)", r"\1\2.\3", text)
    text = re.sub(r"([图表])([0-9A-Z])\.(\d+[a-z]?)/(\d+[a-z]?)", r"\1\2.\3/\2.\4", text)
    return text


def set_run_font(run, size: float | None = None, bold: bool | None = None, font: str | None = None):
    font_name = font or BODY_FONT
    latin_name = LATIN_FONT if font_name == BODY_FONT else font_name
    run.font.name = latin_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_run_font_mixed(
    run,
    *,
    east_asia: str,
    latin: str,
    size: float,
    bold: bool | None = None,
):
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, font_name: str, size: float, bold: bool = False):
    latin_name = LATIN_FONT if font_name == BODY_FONT else font_name
    style.font.name = latin_name
    style.font.size = Pt(size)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attr in ("w:asciiTheme", "w:eastAsiaTheme", "w:hAnsiTheme", "w:cstheme"):
        key = qn(attr)
        if key in r_fonts.attrib:
            del r_fonts.attrib[key]
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), latin_name)
    r_fonts.set(qn("w:hAnsi"), latin_name)
    for tag in ("w:sz", "w:szCs"):
        node = r_pr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            r_pr.append(node)
        node.set(qn("w:val"), str(int(size * 2)))
    bold_node = r_pr.find(qn("w:b"))
    if bold:
        if bold_node is None:
            bold_node = OxmlElement("w:b")
            r_pr.append(bold_node)
        bold_node.set(qn("w:val"), "true")
    elif bold_node is not None:
        bold_node.set(qn("w:val"), "false")
    bold_cs_node = r_pr.find(qn("w:bCs"))
    if bold:
        if bold_cs_node is None:
            bold_cs_node = OxmlElement("w:bCs")
            r_pr.append(bold_cs_node)
        bold_cs_node.set(qn("w:val"), "true")
    elif bold_cs_node is not None:
        bold_cs_node.set(qn("w:val"), "false")
    color = r_pr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        r_pr.append(color)
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        key = qn(attr)
        if key in color.attrib:
            del color.attrib[key]
    color.set(qn("w:val"), "000000")


def set_paragraph_format(paragraph, *, first_indent=True, align=None, line_spacing=1.25, before=0, after=0):
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.first_line_indent = Cm(0.74) if first_indent else Cm(0)
    if align is not None:
        paragraph.alignment = align


def remove_markdown_link(text: str) -> str:
    text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", text)
    text = text.replace("&nbsp;", " ")
    return text


def clean_text(text: str) -> str:
    text = remove_markdown_link(text)
    text = text.replace("\\_", "_")
    text = normalize_figure_table_numbers(text)
    return text.strip()


def add_inline(
    paragraph,
    text: str,
    base_size: float = 12,
    font: str | None = None,
    *,
    code_as_body: bool = False,
):
    """Add a compact subset of Markdown inline formatting."""
    text = clean_text(text)
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, base_size, font=font)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, base_size, bold=True, font=font)
        else:
            run = paragraph.add_run(token[1:-1])
            if code_as_body:
                set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=base_size)
            else:
                set_run_font(run, 10, font=CODE_FONT)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, base_size, font=font)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="BFBFBF"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        spec = kwargs[edge]
        if spec is None:
            element.set(qn("w:val"), "nil")
            continue
        for key, value in spec.items():
            element.set(qn(f"w:{key}"), str(value))


def set_three_line_table_borders(table):
    """Apply a thesis-template-style three-line table: black 0.5pt top/header/bottom."""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge in {"top", "bottom"}:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "4")
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
        else:
            element.set(qn("w:val"), "nil")

    if not table.rows:
        return
    header_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    header_tr_pr.append(tbl_header)
    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            bottom={"val": "single", "sz": "4", "space": "0", "color": "000000"},
        )


def hide_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run._r.append(fld)
    return run


def set_page_numbering(section, fmt: str | None = None, start: int | None = None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    elif pg_num.get(qn("w:fmt")) is not None:
        del pg_num.attrib[qn("w:fmt")]
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    elif pg_num.get(qn("w:start")) is not None:
        del pg_num.attrib[qn("w:start")]


def configure_section(section, header_text: str | None = None, page_num_fmt: str | None = None, start: int | None = None):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.header_distance = Cm(2)
    section.footer_distance = Cm(2)

    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False

    header = section.header.paragraphs[0]
    header.clear()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_text:
        run = header.add_run(SCHOOL_HEADER)
        set_run_font(run, 10.5, font=HEADING_FONT)

    even_header = section.even_page_header.paragraphs[0]
    even_header.clear()
    even_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_text:
        run = even_header.add_run(normalize_heading_text(header_text))
        set_run_font(run, 10.5, font=HEADING_FONT)

    for footer_part in (section.footer, section.even_page_footer):
        footer = footer_part.paragraphs[0]
        footer.clear()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if header_text:
            run = add_field(footer, "PAGE")
            set_run_font_mixed(run, east_asia=BODY_FONT, latin="Arial", size=10.5)

    set_page_numbering(section, page_num_fmt, start)


def configure_styles(doc: Document):
    styles = doc.styles

    set_style_font(styles["Normal"], BODY_FONT, 12)
    pf = styles["Normal"].paragraph_format
    pf.line_spacing = 1.25
    pf.first_line_indent = Cm(0.74)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    set_style_font(styles["Title"], HEADING_FONT, 18, bold=True)
    styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_style_font(styles["Heading 1"], HEADING_FONT, 15, bold=False)
    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 1"].paragraph_format.line_spacing = 1.5
    styles["Heading 1"].paragraph_format.space_before = Pt(0)
    styles["Heading 1"].paragraph_format.space_after = Pt(11)
    styles["Heading 1"].paragraph_format.first_line_indent = None

    set_style_font(styles["Heading 2"], HEADING_FONT, 14, bold=False)
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 2"].paragraph_format.line_spacing = 1.5
    styles["Heading 2"].paragraph_format.space_before = Pt(7.5)
    styles["Heading 2"].paragraph_format.space_after = Pt(0)
    styles["Heading 2"].paragraph_format.first_line_indent = None

    set_style_font(styles["Heading 3"], HEADING_FONT, 12, bold=False)
    styles["Heading 3"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 3"].paragraph_format.line_spacing = 1.5
    styles["Heading 3"].paragraph_format.space_before = Pt(7.5)
    styles["Heading 3"].paragraph_format.space_after = Pt(0)
    styles["Heading 3"].paragraph_format.first_line_indent = None

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        set_style_font(style, BODY_FONT, 12)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_after = Pt(0)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    cells: list[str] = []
    current: list[str] = []
    in_code = False
    for ch in line:
        if ch == "`":
            in_code = not in_code
            current.append(ch)
            continue
        if ch == "|" and not in_code:
            cells.append(clean_text("".join(current)))
            current = []
            continue
        current.append(ch)
    cells.append(clean_text("".join(current)))
    return cells


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*[:\- ]+(\|\s*[:\- ]+)+\|?\s*$", line))


def add_caption_paragraph(doc: Document, text: str, *, kind: str):
    para = doc.add_paragraph()
    set_paragraph_format(para, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.25)
    para.paragraph_format.keep_with_next = kind == "table"
    run = para.add_run(clean_text(text))
    set_run_font_mixed(run, east_asia=HEADING_FONT, latin="Arial", size=10.5, bold=False)
    return para


def add_blank_paragraph(doc: Document):
    para = doc.add_paragraph()
    set_paragraph_format(para, first_indent=False, line_spacing=1.25)
    return para


def add_equation(doc: Document, formula: str, number: str):
    para = doc.add_paragraph()
    set_paragraph_format(para, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.25, before=6, after=6)
    para.paragraph_format.tab_stops.add_tab_stop(Cm(7.5), WD_TAB_ALIGNMENT.CENTER)
    para.paragraph_format.tab_stops.add_tab_stop(Cm(15), WD_TAB_ALIGNMENT.RIGHT)
    run = para.add_run()
    run.add_tab()
    set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=12)
    run = para.add_run(clean_text(formula))
    set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=12)
    run = para.add_run()
    run.add_tab()
    set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=12)
    run = para.add_run(number)
    set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=12)
    return para


def add_table(doc: Document, rows: list[list[str]], caption: str | None = None):
    if not rows:
        return
    if caption:
        add_blank_paragraph(doc)
        add_caption_paragraph(doc, caption, kind="table")
    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    table.style = "Table Grid"
    set_three_line_table_borders(table)

    for r_idx, row in enumerate(normalized):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=80, start=120, bottom=80, end=120)
            set_cell_border(
                cell,
                top=None,
                left=None,
                bottom=None,
                right=None,
                insideH=None,
                insideV=None,
            )
            if r_idx == 0:
                set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": "4", "space": "0", "color": "000000"},
                )
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(value) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(paragraph, first_indent=False, align=paragraph.alignment)
            add_inline(paragraph, value, base_size=10.5)
            for run in paragraph.runs:
                if r_idx == 0:
                    run.bold = True

    add_blank_paragraph(doc)


def add_image(doc: Document, rel_path: str, caption: str | None):
    image_path = (ROOT / rel_path).resolve()
    if not image_path.is_file():
        para = doc.add_paragraph()
        add_inline(para, f"（图文件未找到：{rel_path}）")
        return
    add_blank_paragraph(doc)
    para = doc.add_paragraph()
    set_paragraph_format(para, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.keep_with_next = True
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(14.2))
    if caption:
        add_caption_paragraph(doc, caption, kind="figure")
    add_blank_paragraph(doc)


def add_cover(doc: Document, thesis_title: str):
    configure_section(doc.sections[0], header_text=None)
    doc.sections[0].footer.paragraphs[0].clear()

    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    set_paragraph_format(title, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    run = title.add_run("本科毕业设计（论文）")
    set_run_font(run, 22, bold=True, font=HEADING_FONT)

    for _ in range(4):
        doc.add_paragraph()

    table = doc.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hide_table_borders(table)
    rows = [
        ("题目", thesis_title),
        ("题目续行", "（不续行请删除此行）"),
        ("学院名称", "<填写学院>"),
        ("专业名称", "<填写专业>"),
        ("学生姓名", "<填写姓名>"),
        ("指导教师", "<填写姓名>"),
        ("起止时间", "    年  月  日  至     年  月  日"),
    ]
    for i, (label, value) in enumerate(rows):
        left = table.cell(i, 0)
        right = table.cell(i, 1)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=120, start=120, bottom=120, end=120)
            p = cell.paragraphs[0]
            set_paragraph_format(p, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5)
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_inline(lp, label, base_size=14, font=HEADING_FONT)
        rp = right.paragraphs[0]
        add_inline(rp, value, base_size=14)


def normalize_heading_text(text: str) -> str:
    text = clean_text(text)
    if text == "摘要":
        return "摘    要"
    if text == "目录":
        return "目    录"
    text = re.sub(r"^(第\d+章)\s*", r"\1  ", text)
    text = re.sub(r"^(\d+(?:\.\d+)+)\s+", r"\1  ", text)
    return text


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    add_inline(paragraph, normalize_heading_text(text), base_size={1: 15, 2: 14, 3: 12}.get(min(level, 3), 12), font=HEADING_FONT)
    return paragraph


def add_toc(doc: Document):
    heading = add_heading(doc, "目    录", 1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    set_paragraph_format(p, first_indent=False)
    add_field(p, r'TOC \o "1-3" \h \z \u')


def update_fields_with_word(path: Path):
    """Update TOC and page fields when Microsoft Word is available."""
    if sys.platform != "win32":
        return
    script = f"""
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
  $doc = $word.Documents.Open('{str(path)}')
  foreach ($toc in $doc.TablesOfContents) {{ $toc.Update() }}
  $doc.Fields.Update() | Out-Null
  $doc.Save()
  $doc.Close()
}} finally {{
  $word.Quit()
}}
"""
    try:
        subprocess.run(["powershell", "-NoProfile", "-Command", script], check=True, timeout=120)
    except Exception as exc:
        print(f"WARN: could not update Word fields automatically: {exc}")


def finalize_docx_styles(path: Path):
    doc = Document(str(path))
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    doc.save(str(path))


def is_appendix_caption(text: str) -> bool:
    compact = re.sub(r"\s+", "", text.strip())
    return bool(re.match(r"^[图表][A-Z0-9]\.\d+[a-z]?", compact))


def enforce_appendix_body_fonts(path: Path):
    """Write appendix body fonts directly so the submitted DOCX is inspectable."""
    doc = Document(str(path))
    in_appendix = False
    for para in doc.paragraphs:
        text = para.text.strip()
        style_name = para.style.name if para.style is not None else ""
        if style_name == "Heading 1" and re.match(r"^附录[A-Z]", text):
            in_appendix = True
            continue
        if not in_appendix:
            continue
        if style_name in {"Heading 1", "Heading 2", "Heading 3"}:
            continue
        if is_appendix_caption(text):
            continue
        for run in para.runs:
            if run.text:
                set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=12)
    doc.save(str(path))


def extract_title(lines: list[str]) -> str:
    for line in lines:
        match = re.match(r"^#\s+(.+)$", line.strip())
        if match:
            return clean_text(match.group(1))
    return "基于 Flaky 分析与自适应执行策略的自动化测试平台设计与实现"


def convert():
    lines = SRC.read_text(encoding="utf-8").splitlines()
    thesis_title = extract_title(lines)

    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    doc.core_properties.title = thesis_title
    doc.core_properties.author = ""
    add_cover(doc, thesis_title)

    current_part = "cover"
    main_started = False
    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[str] = []
    pending_table_caption: str | None = None
    skip_manual_toc = False

    def new_section(header_text: str | None, fmt: str | None = None, start: int | None = None):
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, header_text=header_text, page_num_fmt=fmt, start=start)
        return section

    new_section("摘    要", fmt="upperRoman", start=1)

    def flush_table():
        nonlocal table_buffer, pending_table_caption
        if not table_buffer:
            return
        rows = []
        for row_line in table_buffer:
            if is_table_separator(row_line):
                continue
            rows.append(parse_table_row(row_line))
        add_table(doc, rows, caption=pending_table_caption)
        table_buffer = []
        pending_table_caption = None

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        para = doc.add_paragraph()
        if current_part == "appendix":
            set_paragraph_format(para, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.25)
        else:
            set_paragraph_format(para, first_indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.0)
        run = para.add_run("\n".join(code_buffer))
        if current_part == "appendix":
            set_run_font_mixed(run, east_asia=BODY_FONT, latin=LATIN_FONT, size=12)
        else:
            set_run_font(run, 9, font=CODE_FONT)
        code_buffer = []

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            continue

        if stripped == f"# {thesis_title}":
            continue

        if stripped.startswith(">"):
            continue

        if stripped.startswith("```"):
            flush_table()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if skip_manual_toc:
            manual_toc_end = re.sub(r"\s+", "", stripped)
            if manual_toc_end.startswith("#前言") or manual_toc_end.startswith("#引言") or stripped.startswith("# 第1章"):
                skip_manual_toc = False
            else:
                continue

        table_caption_match = re.match(r"^\*\*(表\s*[0-9A-Z][.-]\d+[a-z]?[^*]*)\*\*\s*(.*)$", stripped)
        if table_caption_match:
            flush_table()
            pending_table_caption = (
                table_caption_match.group(1)
                + (f" {table_caption_match.group(2).strip()}" if table_caption_match.group(2).strip() else "")
            ).strip()
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            table_buffer.append(stripped)
            continue
        flush_table()

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if image_match:
            add_image(doc, image_match.group(2), image_match.group(1))
            continue

        equation_match = re.match(r"^\$\$\s*(.+?)\s*\\tag\{([^}]+)\}\s*\$\$$", stripped)
        if equation_match:
            add_equation(doc, equation_match.group(1), f"({equation_match.group(2)})")
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            raw_level = len(heading_match.group(1))
            heading = clean_text(heading_match.group(2))

            if heading == "摘要":
                current_part = "abstract_cn"
                add_heading(doc, "摘    要", 1)
                continue
            if heading == "ABSTRACT":
                current_part = "abstract_en"
                new_section("ABSTRACT", fmt="upperRoman")
                title_para = doc.add_paragraph()
                set_paragraph_format(title_para, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
                title_para.paragraph_format.line_spacing = 1.25
                title_para.paragraph_format.space_before = Pt(0)
                title_para.paragraph_format.space_after = Pt(0)
                add_inline(title_para, ENGLISH_TITLE, base_size=15, font=LATIN_FONT)
                doc.add_paragraph()
                add_heading(doc, "ABSTRACT", 1)
                continue
            if heading == "目录":
                current_part = "toc"
                new_section("目    录", fmt="upperRoman")
                add_toc(doc)
                skip_manual_toc = True
                continue

            promoted = False
            if heading in {"致谢", "参考文献"} or heading.startswith("附录"):
                promoted = True

            normalized_heading = re.sub(r"\s+", "", heading)

            if raw_level == 1 or promoted:
                if normalized_heading in {"前言", "引言"}:
                    current_part = "body"
                    if not main_started:
                        new_section(heading, fmt="decimal", start=1)
                        main_started = True
                    else:
                        new_section(heading, fmt="decimal")
                elif heading.startswith("第") and "章" in heading:
                    current_part = "body"
                    if not main_started:
                        new_section(SCHOOL_HEADER, fmt="decimal", start=1)
                        main_started = True
                    else:
                        new_section(heading, fmt="decimal")
                elif heading.startswith("附录"):
                    current_part = "appendix"
                    new_section(heading, fmt="decimal")
                else:
                    new_section(heading, fmt="decimal")
                add_heading(doc, heading, 1)
                continue

            add_heading(doc, heading, min(raw_level, 3))
            continue

        if re.match(r"^[-*+]\s+", stripped):
            item = re.sub(r"^[-*+]\s+", "", stripped)
            para = doc.add_paragraph(style="List Bullet")
            set_paragraph_format(para, first_indent=False)
            add_inline(
                para,
                item,
                base_size=12,
                font=LATIN_FONT if current_part == "abstract_en" else None,
                code_as_body=current_part == "appendix",
            )
            continue

        if re.match(r"^\d+\.\s+", stripped):
            item = re.sub(r"^\d+\.\s+", "", stripped)
            para = doc.add_paragraph(style="List Number")
            set_paragraph_format(para, first_indent=False)
            add_inline(
                para,
                item,
                base_size=12,
                font=LATIN_FONT if current_part == "abstract_en" else None,
                code_as_body=current_part == "appendix",
            )
            continue

        para = doc.add_paragraph()
        if current_part == "abstract_en" and stripped.startswith("**Title:**"):
            set_paragraph_format(para, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            add_inline(para, stripped.replace("**Title:**", "").strip(), base_size=15, font=LATIN_FONT)
        else:
            set_paragraph_format(para)
            add_inline(
                para,
                stripped,
                base_size=12,
                font=LATIN_FONT if current_part == "abstract_en" else None,
                code_as_body=current_part == "appendix",
            )

    flush_table()
    flush_code()
    doc.save(OUT_DOCX)
    update_fields_with_word(OUT_DOCX.resolve())
    finalize_docx_styles(OUT_DOCX)
    enforce_appendix_body_fonts(OUT_DOCX)
    print(f"OK: {SRC.name} -> {OUT_DOCX.name}")


if __name__ == "__main__":
    convert()
